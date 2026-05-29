from __future__ import annotations

import random
from datetime import date, timedelta
from pathlib import Path

import fitz
from docx import Document


OUTPUT_DIR = Path("data/sample")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

random.seed(42)

FACULTY_PROGRAMS = {
    "Engineering": [
        "Industrial Engineering",
        "Systems Engineering",
        "Civil Engineering",
        "Environmental Engineering",
    ],
    "Business": [
        "Business Administration",
        "Marketing",
        "Finance",
    ],
    "Communication": [
        "Communication",
        "Advertising",
    ],
}

REQUEST_TYPES = ["APC", "IDIC", "Ethics Review", "Patent", "Research Project"]
DOCUMENT_TYPES = ["Application Form", "Committee Report", "Review Memo", "Decision Letter"]
STATUSES = ["Approved", "Approved with observations", "Pending", "Rejected", "Under review"]

OBSERVATIONS = [
    "No major observations.",
    "Missing informed consent form.",
    "Methodology section requires clarification.",
    "Budget justification requires additional detail.",
    "Advisor signature is missing.",
    "Research scope should be narrowed.",
    "Ethical risk assessment requires more detail.",
    "Document is incomplete and needs resubmission.",
]

RESULTS = [
    "Approved",
    "Conditionally approved",
    "Pending additional documentation",
    "Rejected",
    "Under committee review",
]

TITLES = [
    "Automation of Administrative Case Monitoring",
    "Digital Transformation in University Services",
    "Data Quality Assessment in Academic Workflows",
    "Optimization of Internal Review Processes",
    "Use of Analytics for Research Management",
    "Process Improvement in Document-Based Workflows",
    "Information Systems for Academic Decision-Making",
    "Operational Indicators for Committee Management",
]


def clean_existing_samples() -> None:
    for pattern in ["case_*.txt", "case_*.docx", "case_*.pdf"]:
        for file_path in OUTPUT_DIR.glob(pattern):
            file_path.unlink()


def build_case(case_number: int) -> dict:
    faculty = random.choice(list(FACULTY_PROGRAMS.keys()))
    program = random.choice(FACULTY_PROGRAMS[faculty])

    submission_date = date(2025, 1, 1) + timedelta(days=random.randint(0, 300))
    review_date = submission_date + timedelta(days=random.randint(3, 20))

    status = random.choice(STATUSES)

    if status == "Pending":
        decision_date = ""
        result = "Pending additional documentation"
    elif status == "Under review":
        decision_date = ""
        result = "Under committee review"
    else:
        decision_date_value = review_date + timedelta(days=random.randint(2, 30))
        decision_date = decision_date_value.isoformat()
        result = random.choice(RESULTS)

    # Introduce a few realistic quality issues.
    advisor = f"Advisor_{random.randint(1, 12):03d}"
    coauthors = ", ".join(
        [f"Coauthor_{random.randint(1, 20):03d}" for _ in range(random.randint(0, 3))]
    )

    if case_number in {7, 18}:
        program = ""

    if case_number in {11, 22}:
        advisor = ""

    if case_number == 15:
        decision_date = (submission_date - timedelta(days=5)).isoformat()

    if case_number == 24:
        decision_date = "2025-13-40"

    return {
        "Case ID": f"CASE-2025-{case_number:03d}",
        "Document Type": random.choice(DOCUMENT_TYPES),
        "Request Type": random.choice(REQUEST_TYPES),
        "Title": random.choice(TITLES),
        "Applicant": f"Applicant_{case_number:03d}",
        "Coauthors": coauthors,
        "Advisor": advisor,
        "Faculty": faculty,
        "Program": program,
        "Submission Date": submission_date.isoformat(),
        "Review Date": review_date.isoformat(),
        "Decision Date": decision_date,
        "Status": status,
        "Result": result,
        "Observations": random.choice(OBSERVATIONS),
    }


def case_to_text(case: dict) -> str:
    return "\n".join(f"{key}: {value}" for key, value in case.items())


def write_txt(file_path: Path, content: str) -> None:
    file_path.write_text(content, encoding="utf-8")


def write_docx(file_path: Path, content: str) -> None:
    doc = Document()
    doc.add_heading("Administrative Case Document", level=1)
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(file_path)


def write_pdf(file_path: Path, content: str) -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    y = 72

    page.insert_text((72, y), "Administrative Case Document", fontsize=14)
    y += 28

    for line in content.splitlines():
        if y > 760:
            page = pdf.new_page()
            y = 72
        page.insert_text((72, y), line, fontsize=10)
        y += 16

    pdf.save(file_path)
    pdf.close()


def main(total_cases: int = 30) -> None:
    clean_existing_samples()

    for case_number in range(1, total_cases + 1):
        case = build_case(case_number)
        content = case_to_text(case)

        if case_number % 3 == 1:
            file_path = OUTPUT_DIR / f"case_2025_{case_number:03d}.txt"
            write_txt(file_path, content)
        elif case_number % 3 == 2:
            file_path = OUTPUT_DIR / f"case_2025_{case_number:03d}.docx"
            write_docx(file_path, content)
        else:
            file_path = OUTPUT_DIR / f"case_2025_{case_number:03d}.pdf"
            write_pdf(file_path, content)

    print(f"Generated {total_cases} synthetic documents in {OUTPUT_DIR.resolve()}")


if __name__ == "__main__":
    main()